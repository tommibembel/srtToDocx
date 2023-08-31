import io
import re
import os
import sys
import datetime
import docx
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
import xlsxwriter
import configparser

config_dict = {"Main": {"UIScale": "100%",
                        "Appearance": "System",
                        "Color": "blue",
                        "OutputFormat": "docx"},
               "Formats": {"docx": ("MS Word", "*.docx"),
                          "xlsx": ("MS Excel", "*.xlsx")}}



def check_configfile_exists(cf):
    if os.path.isfile(cf):
        config = configparser.RawConfigParser()
        config.read(cf)
        for i in config_dict.keys():
            if config.has_section(i):
                return config
        else:
            config = create_cofig_file(cf)
            return config
    else:
        config = create_cofig_file(cf)
        return config


def create_cofig_file(cf):
    config = configparser.RawConfigParser()
    config["Main"] = config_dict["Main"]
    config["Formats"] = config_dict["Formats"]
    with open(cf, "w") as configfile:
        config.write(configfile)
    return config


def write_config(cf, conf):
    with open(cf, "w") as configfile:
        conf.write(configfile)


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def parse_srt(inputfile):
    # Check if file ends with srt. so no wrong filetype is chosen.
    # Then we parse the srt file for indexes, timecodes and subtitles, separated by newlines and in case
    # of subtitles separated by a blank line
    if inputfile.endswith('.srt'):
        f = io.open(inputfile, encoding="utf-8-sig", mode="r")
        time_re = re.compile(r'\d{2}:\d{2}:\d{2},\d{3} --> \d{2}:\d{2}:\d{2},\d{3}')
        result = {}
        key = ""
        subtitle = ""
        for line in f.readlines():
            if line.strip().isdigit():
                key = line

            elif re.search(time_re, line):
                times = line.strip().split(' --> ')
                dt_start = datetime.datetime.strptime(times[0], "%H:%M:%S,%f")
                dt_end = datetime.datetime.strptime(times[1], "%H:%M:%S,%f")
                dt_dur = dt_end - dt_start

            elif line in ["\n", "\r\n"] and not subtitle == "" and not key == "":
                result[int(key)] = (
                dt_start.strftime("%H:%M:%S,%f")[:-3], dt_end.strftime("%H:%M:%S,%f")[:-3], str(dt_dur)[:-3], subtitle)
                subtitle = ""
                key = ""

            elif line in ["\n", "\r\n"]:
                print("Error in line")

            else:
                subtitle += line

        # The last subtitle may not contain a blank line at the end, so we have to check if the subtitle and index
        # variable are not empty and add it to the result dictionary
        if subtitle != "" and key != "":
            result[int(key)] = (
                dt_start.strftime("%H:%M:%S,%f")[:-3], dt_end.strftime("%H:%M:%S,%f")[:-3], str(dt_dur)[:-3], subtitle)

        return result


def write_docx(outputfile, result):
    # Generating the docx file and adding a table to it.
    document = docx.Document()
    table = document.add_table(rows=1, cols=5, style='Table Grid')
    table.allow_autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    # Adding headers to the table
    tableheaders = table.rows[0].cells
    tableheaders[0].text = '#'
    tableheaders[1].text = 'Start Time'
    tableheaders[2].text = 'End Time'
    tableheaders[3].text = 'Duration'
    tableheaders[4].text = 'Text'

    # Set fontcolor to black and bg color of to white for table headers
    for i in range(len(tableheaders)):
        tableheaders[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        bgc = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(bgc)

    # Adding our result dictionary to the table row by row. If there is a tag for italic words (<i> / </i>) in the
    # text, we need to rebuild this paragraph to add the italic.
    for index, (startTime, endTime, duration, subtitle) in sorted(result.items()):
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        row_cells[1].text = str(startTime)
        row_cells[2].text = str(endTime)
        row_cells[3].text = str(duration)
        if subtitle.find("<i>") >= 0:
            p = row_cells[4].paragraphs[0].insert_paragraph_before(text=None, style=None)
            italic_start_indexes = [i.start() for i in re.finditer('<i>', subtitle)]
            italic_end_indexes = [i.start() for i in re.finditer('</i>', subtitle)]
            if len(italic_start_indexes) != len(italic_end_indexes):
                return 1
            pointer = 0
            for m in re.finditer('<i>(.*?)</i>', subtitle):
                before = subtitle[pointer:m.start()]
                italic_text = m.group(1)
                pointer = m.end()
                p.add_run(before)
                run = p.add_run(italic_text)
                run.italic = True
            after = subtitle[pointer:]
            p.add_run(after.removesuffix("\n"))
            """
            for line in subtitle.removesuffix("\n").splitlines(True):
                if line.find("<i>") >= 0:
                    index1 = line.index("<i>")
                    index2 = line.index("</i>")
                    before = line[0:index1]
                    after = line[index2 + 4:]
                    italic_text = line[index1 + 3:index2]
                    p.add_run(before)
                    run = p.add_run(italic_text)
                    run.italic = True
                    p.add_run(after)
            """
        else:
            row_cells[4].text = str(subtitle)

    # Formatting the columns. I did not find a better solution than iterate over the whole document and set
    # each cell individually to a certain width.
    for cell in table.columns[0].cells:
        cell.width = 0.5 * 914400

    for cell in table.columns[4].cells:
        cell.width = 4 * 914400

    # Try to save the document. If it is reconverted and already opened,
    # or if we do not have any write permissions, we have to catch the exception.
    try:
        document.save(outputfile)
        return 0
    except Exception as e:
        return e


def write_xlsx(outputfile, result):
    workbook = xlsxwriter.Workbook(outputfile)
    textwrap = workbook.add_format({'text_wrap': True,})
    italic = workbook.add_format({'italic': True})
    align_top = workbook.add_format({'valign': 'top', 'text_wrap': True})

    worksheet = workbook.add_worksheet()
    worksheet.set_column(0, 4, cell_format=align_top)
    row, col = 1, 0
    tableheaders = ['#', 'Start Time', 'End Time', 'Duration', 'Text']
    for i, h in enumerate(tableheaders):
        worksheet.write(0, i, h)

    for index, (startTime, endTime, duration, subtitle) in sorted(result.items()):
        worksheet.write_string(row, 0, str(index))
        worksheet.write_string(row, 1, str(startTime))
        worksheet.write_string(row, 2, str(endTime))
        worksheet.write_string(row, 3, str(duration))

        if subtitle.find("<i>") >= 0:
            italic_start_indexes = [i.start() for i in re.finditer('<i>', subtitle)]
            italic_end_indexes = [i.start() for i in re.finditer('</i>', subtitle)]
            if len(italic_start_indexes) != len(italic_end_indexes):
                return 1
            richtext = []
            pointer = 0
            for m in re.finditer('<i>(.*?)</i>', subtitle):
                before = subtitle[pointer:m.start()]
                italic_text = m.group(1)
                pointer = m.end()
                if before != "":
                    richtext.append(before)
                richtext.append(italic)
                richtext.append(italic_text)
            after = subtitle[pointer:]
            if after.removesuffix("\n") != "":
                richtext.append(after.removesuffix("\n"))
            if len(richtext) == 2:
                worksheet.write_string(row, 4, str(subtitle).removesuffix("\n")
                                       .replace("<i>","").replace("</i>",""), italic)
            else:
                worksheet.write_rich_string(row, 4, *richtext)
        else:
            worksheet.write_string(row, 4, str(subtitle).removesuffix("\n"))

        row += 1

    worksheet.autofit()
    workbook.close()
    return 0
